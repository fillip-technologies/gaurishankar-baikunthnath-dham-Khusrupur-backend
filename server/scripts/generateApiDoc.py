# Generates docs/API_Documentation.docx — a full per-endpoint API reference.
from datetime import date
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MAROON = RGBColor(0xB7, 0x1C, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x55, 0x55, 0x55)
METHOD_COLORS = {
    "GET": RGBColor(0x2E, 0x7D, 0x32),
    "POST": RGBColor(0x15, 0x65, 0xC0),
    "PATCH": RGBColor(0xE6, 0x5A, 0x00),
    "DELETE": RGBColor(0xC6, 0x28, 0x28),
}

# Each endpoint:
#   method, path, access, content_type, rate_limit, notes,
#   params:  [ (name, type, rules) ]
#   query:   [ (name, type, rules) ]
#   body:    [ (field, type, rules, required) ]
#   responses: [ (status, description) ]
MODULES = [
    ("Authentication", "/api/v1/auth", [
        dict(method="POST", path="/login", access="Public", content_type="application/json",
             rate_limit="loginLimiter",
             notes="From a trusted device returns a session; from a new device sends an email OTP and returns a challenge cookie instead.",
             body=[("email", "string", "valid email", "yes"),
                   ("password", "string", "6-100 chars", "yes")],
             responses=[("200 (trusted device)", "{ data: { user }, message: 'Login successful' } + sets accessToken, refreshToken, deviceId cookies."),
                        ("200 (new device)", "{ data: { requiresOtp: true, email }, message: '...OTP sent' } + sets loginChallenge cookie."),
                        ("401", "Invalid email or password.")]),
        dict(method="POST", path="/verify_login_otp", access="Public", content_type="application/json",
             rate_limit="otpLimiter",
             notes="Completes a new-device login. Requires the loginChallenge cookie from /login.",
             body=[("otp", "string", "exactly 6 digits", "yes")],
             responses=[("200", "{ data: { user }, message: 'Login successful' } + sets auth cookies, clears challenge."),
                        ("401", "Invalid/expired OTP or too many attempts.")]),
        dict(method="POST", path="/create_admin", access="Auth + Session (superadmin)", content_type="application/json",
             rate_limit="apiRateLimiter",
             notes="Superadmin only. Emails the new admin their credentials; rolls back if the email fails.",
             body=[("fullname", "string", "3-50 chars", "yes"),
                   ("mobile_number", "string", "10-15 chars", "yes"),
                   ("email", "string", "valid email", "yes"),
                   ("password", "string", "6-100 chars", "yes")],
             responses=[("201", "{ data: null, message: 'Admin created successfully. Login credentials have been emailed.' }"),
                        ("403", "Caller is not a superadmin."),
                        ("409", "Email already in use.")]),
        dict(method="POST", path="/refresh_token", access="Public (refresh cookie)", content_type="-",
             rate_limit="apiRateLimiter",
             notes="Reads the refreshToken cookie; rotates both tokens.",
             responses=[("200", "{ message: 'token updated' } + new accessToken & refreshToken cookies."),
                        ("401", "Missing/expired/invalid refresh token.")]),
        dict(method="PATCH", path="/update_password", access="Auth + Session", content_type="application/json",
             rate_limit="apiRateLimiter",
             body=[("current_password", "string", "6-100 chars", "yes"),
                   ("newpassword", "string", "6-100 chars", "yes")],
             responses=[("200", "{ data: null, message: 'Password succesfully changed' }"),
                        ("401", "Current password incorrect.")]),
        dict(method="GET", path="/admins", access="Auth + Session (superadmin)", content_type="-",
             rate_limit="apiRateLimiter",
             notes="Superadmin only.",
             responses=[("200", "{ data: [ admin, ... ] (role 'admin'), message }"),
                        ("403", "Caller is not a superadmin.")]),
        dict(method="GET", path="/profile/:id", access="Auth", content_type="-",
             rate_limit="apiRateLimiter",
             notes="Allowed for the admin themselves or a superadmin.",
             params=[("id", "string", "admin _id")],
             responses=[("200", "{ data: admin (sensitive fields stripped), message }"),
                        ("401", "Not the owner and not a superadmin.")]),
        dict(method="POST", path="/logout", access="Auth", content_type="-", rate_limit="-",
             responses=[("200", "{ data: null, message: 'User sucussfully logged out' } + clears auth cookies.")]),
        dict(method="DELETE", path="/remove_admin", access="Auth + Session (superadmin)", content_type="application/json",
             rate_limit="apiRateLimiter",
             notes="Superadmin only. Cannot delete your own account or another superadmin.",
             body=[("adminEmail", "string", "valid email", "yes"),
                   ("superAdminPassword", "string", "6-100 chars", "yes")],
             responses=[("200", "{ success: true, message: 'Admin deleted successfully' }"),
                        ("401/403/404", "Wrong password / not superadmin / admin not found.")]),
    ]),
    ("Addresses", "/api/v1/addresses", [
        dict(method="POST", path="/", access="Auth + Session", content_type="application/json",
             rate_limit="apiRateLimiter",
             notes="Sets the current admin's address. All fields optional (partial update).",
             body=[("houseName", "string", "non-empty", "no"),
                   ("locality", "string", "non-empty", "no"),
                   ("district", "string", "non-empty", "no"),
                   ("state", "string", "non-empty", "no"),
                   ("zipcode", "string", "exactly 6 digits", "no"),
                   ("country", "string", "non-empty", "no")],
             responses=[("200", "{ data: address, message: 'Address updated succesfully' }")]),
        dict(method="PATCH", path="/", access="Auth + Session", content_type="application/json",
             rate_limit="apiRateLimiter",
             notes="Same handler as POST / (create or update).",
             body=[("(same fields as POST /)", "", "", "")],
             responses=[("200", "{ data: address, message: 'Address updated succesfully' }")]),
        dict(method="GET", path="/", access="Auth", content_type="-", rate_limit="apiRateLimiter",
             responses=[("200", "{ data: address, message: 'Address fetched succesfully' }")]),
    ]),
    ("Media", "/api/v1/media", [
        dict(method="POST", path="/galleryUpload", access="Auth", content_type="multipart/form-data",
             rate_limit="apiRateLimiter",
             notes="Image is uploaded to Cloudinary.",
             body=[("file", "file", "image, required (multipart)", "yes"),
                   ("dataType", "enum", "photos | wallpaper | videos", "yes"),
                   ("title", "string", ">= 3 chars", "no"),
                   ("occasion", "string", "3-50 chars", "no")],
             responses=[("201", "{ data: galleryItem, message: 'Uploaded to gallery' }"),
                        ("400", "No file uploaded.")]),
        dict(method="DELETE", path="/gallery/:id", access="Auth", content_type="-",
             rate_limit="apiRateLimiter",
             params=[("id", "string", "gallery item _id")],
             responses=[("200", "{ data: null, message }  (also removed from Cloudinary)")]),
        dict(method="GET", path="/gallery", access="Public", content_type="-",
             rate_limit="apiRateLimiter",
             notes="Paginated, 20 per page.",
             query=[("page", "number", "integer >= 1, required"),
                    ("dataType", "enum", "photos | wallpaper | videos, required")],
             responses=[("200", "{ data: [ galleryItem, ... ], message }")]),
        dict(method="POST", path="/mediaCoverage", access="Auth", content_type="multipart/form-data",
             rate_limit="apiRateLimiter",
             body=[("file", "file", "required (multipart)", "yes"),
                   ("title", "string", "3-150 chars", "no"),
                   ("description", "string", "3-2000 chars", "no")],
             responses=[("201", "{ data: media, message: 'Media created successfully' }"),
                        ("400", "No file uploaded.")]),
        dict(method="DELETE", path="/mediaCoverage", access="Auth + Session", content_type="-",
             rate_limit="apiRateLimiter",
             query=[("id", "string", "media _id, required")],
             responses=[("200", "{ data: response, message: 'Data deleted Succesfully' }"),
                        ("400", "Missing id.")]),
        dict(method="GET", path="/mediaCoverage/:id", access="Public", content_type="-",
             rate_limit="apiRateLimiter",
             params=[("id", "string", "media _id")],
             responses=[("200", "{ data: media, message }"),
                        ("404", "No media found.")]),
        dict(method="GET", path="/mediaCoverage", access="Public", content_type="-",
             rate_limit="apiRateLimiter",
             notes="Newest first.",
             responses=[("200", "{ data: [ media, ... ], message }"),
                        ("404", "No media found.")]),
    ]),
    ("Payments", "/api/v1/payments", [
        dict(method="POST", path="/webhook", access="Public (Razorpay signature)", content_type="application/json (raw)",
             rate_limit="-",
             notes="Server-to-server callback. Authenticity verified via the x-razorpay-signature header (HMAC over the raw body). Declared first so it is not shadowed.",
             responses=[("200", "{ received: true } — acknowledges and reconciles the payment status.")]),
        dict(method="POST", path="/order", access="Public", content_type="application/json",
             rate_limit="apiRateLimiter",
             notes="Standalone checkout. Domain flows (e.g. prasad booking) call the service directly rather than this route.",
             body=[("payer.name", "string", ">= 2 chars", "yes"),
                   ("payer.email", "string", "valid email", "yes"),
                   ("payer.phone", "string", "10-14 digits", "yes"),
                   ("amount", "number", "integer, paise, >= 100 (Rs.1)", "yes"),
                   ("currency", "string", "3-letter ISO, default INR", "no"),
                   ("purpose", "enum", "booking | donation | general, default general", "no"),
                   ("reference.model", "string", "non-empty (with reference.id)", "no"),
                   ("reference.id", "string", "24-hex ObjectId", "no"),
                   ("notes", "object", "string->string map", "no")],
             responses=[("201", "{ data: { key, orderId, amount, currency, paymentId, payer }, message: 'Order created' }")]),
        dict(method="POST", path="/verify", access="Public", content_type="application/json",
             rate_limit="apiRateLimiter",
             body=[("razorpayOrderId", "string", "required", "yes"),
                   ("razorpayPaymentId", "string", "required", "yes"),
                   ("razorpaySignature", "string", "required", "yes")],
             responses=[("200", "{ data: { paymentId, status, razorpayOrderId, razorpayPaymentId }, message: 'Payment verified' }"),
                        ("400", "Invalid signature.")]),
        dict(method="GET", path="/all", access="Auth + Session (superadmin | accountant)", content_type="-",
             rate_limit="apiRateLimiter",
             notes="Restricted to superadmin/accountant (enforced in controller).",
             responses=[("200", "{ data: [ payment, ... ], message }"),
                        ("401", "Role not permitted.")]),
        dict(method="GET", path="/:id", access="Auth + Session (superadmin | accountant)", content_type="-",
             rate_limit="apiRateLimiter",
             params=[("id", "string", "payment _id")],
             responses=[("200", "{ data: payment, message: 'Payment fetched' }"),
                        ("401", "Role not permitted."), ("404", "Payment not found.")]),
    ]),
    ("Prasad & Bookings", "/api/v1/prasad", [
        dict(method="POST", path="/add", access="Auth + Session", content_type="multipart/form-data",
             rate_limit="apiRateLimiter",
             notes="Admin catalogue. Image uploaded to Cloudinary.",
             body=[("file", "file", "image (multipart)", "no"),
                   ("prasadName", "string", "3-100 chars", "yes"),
                   ("pricePerKg", "number", "positive", "yes"),
                   ("description", "string", "3-500 chars", "no")],
             responses=[("200", "{ data: prasad, message: 'Prasad added!' }")]),
        dict(method="DELETE", path="/remove/:id", access="Auth + Session", content_type="-",
             rate_limit="apiRateLimiter",
             params=[("id", "string", "prasad _id")],
             responses=[("200", "{ data, message: 'Deleted succesfully' }")]),
        dict(method="GET", path="/prasads", access="Public", content_type="-", rate_limit="-",
             responses=[("200", "{ data: [ prasad, ... ], message: 'All prasad' }"),
                        ("204", "No prasad listed.")]),
        dict(method="GET", path="/bookings", access="Auth + Session", content_type="-",
             rate_limit="apiRateLimiter",
             notes="Admin dashboard list, newest first, with prasad + payer joined in.",
             query=[("status", "enum", "pending | confirmed | failed (optional)"),
                    ("from", "date", "YYYY-MM-DD, inclusive (optional)"),
                    ("to", "date", "YYYY-MM-DD, inclusive (optional)"),
                    ("page", "number", "default 1"),
                    ("limit", "number", "default 20, max 100")],
             responses=[("200", "{ data: { bookings: [...], pagination: { total, page, limit, totalPages } }, message }")]),
        dict(method="POST", path="/book", access="Public", content_type="application/json",
             rate_limit="apiRateLimiter",
             notes="Server prices the order from the catalogue; the client never sends the amount.",
             body=[("prasadId", "string", "24-hex ObjectId", "yes"),
                   ("quantity", "number", "positive", "yes"),
                   ("payer.name", "string", "3-50 chars", "yes"),
                   ("payer.email", "string", "valid email", "yes"),
                   ("payer.phone", "string", "10-14 digits", "yes")],
             responses=[("201", "{ data: { bookingId, prasadName, quantity, key, orderId, amount, currency, paymentId, payer }, message: 'Prasad booking created' }"),
                        ("404", "Prasad not found.")]),
        dict(method="POST", path="/verify", access="Public", content_type="application/json",
             rate_limit="apiRateLimiter",
             notes="On success, emails the devotee a receipt with a downloadable PDF invoice.",
             body=[("razorpayOrderId", "string", "required", "yes"),
                   ("razorpayPaymentId", "string", "required", "yes"),
                   ("razorpaySignature", "string", "required", "yes")],
             responses=[("200", "{ data: { payment, booking }, message: 'Payment verified' }"),
                        ("400", "Invalid signature.")]),
    ]),
]


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def field_table(doc, title, headers, rows, widths):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "8A1414")
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
            if i == 0:
                run.bold = True
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# Title page
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Shree Gaurishankar Baikunthdham Temple")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = MAROON
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run("Backend API Documentation")
sr.font.size = Pt(14)
sr.font.color.rgb = GREY
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Base URL: /api/v1     |     Generated: {date.today().isoformat()}").italic = True

# Conventions
doc.add_paragraph()
hp = doc.add_paragraph().add_run("Conventions")
hp.bold = True
hp.font.size = Pt(13)
hp.font.color.rgb = MAROON
for label, desc in [
    ("Response envelope", "All JSON responses use { statusCode, data, message, success } (ApiResponse)."),
    ("Auth (cookies)", "Access/refresh tokens are httpOnly cookies set on login; send them automatically with credentials."),
    ("Public", "no authentication."),
    ("Auth", "valid access token (authenticate)."),
    ("Auth + Session", "access token + active session (authenticate + requireValidSession)."),
    ("(superadmin) / (accountant)", "role additionally enforced inside the controller."),
]:
    p = doc.add_paragraph(style="List Bullet")
    b = p.add_run(f"{label}: ")
    b.bold = True
    p.add_run(desc)

BODY_W = [Inches(1.7), Inches(0.9), Inches(3.2), Inches(0.9)]
PQ_W = [Inches(1.7), Inches(0.9), Inches(4.1)]

for name, prefix, endpoints in MODULES:
    doc.add_page_break()
    hp = doc.add_paragraph()
    hr = hp.add_run(f"{name}")
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = MAROON
    doc.add_paragraph().add_run(f"Mounted at {prefix}").italic = True

    for ep in endpoints:
        doc.add_paragraph()
        head = doc.add_paragraph()
        mrun = head.add_run(ep["method"] + "  ")
        mrun.bold = True
        mrun.font.size = Pt(13)
        mrun.font.color.rgb = METHOD_COLORS.get(ep["method"], RGBColor(0, 0, 0))
        prun = head.add_run(prefix + ep["path"])
        prun.bold = True
        prun.font.size = Pt(13)

        meta = doc.add_paragraph()
        meta_txt = f"Access: {ep['access']}    |    Content-Type: {ep['content_type']}    |    Rate limit: {ep['rate_limit']}"
        mr = meta.add_run(meta_txt)
        mr.font.size = Pt(9)
        mr.font.color.rgb = GREY

        if ep.get("notes"):
            n = doc.add_paragraph()
            nr = n.add_run("Note: " + ep["notes"])
            nr.italic = True
            nr.font.size = Pt(9)

        if ep.get("params"):
            field_table(doc, "Path parameters", ["Param", "Type", "Description"], ep["params"], PQ_W)
        if ep.get("query"):
            field_table(doc, "Query parameters", ["Param", "Type", "Rules"], ep["query"], PQ_W)
        if ep.get("body"):
            field_table(doc, "Request body", ["Field", "Type", "Rules", "Required"], ep["body"], BODY_W)

        rp = doc.add_paragraph()
        rpr = rp.add_run("Responses")
        rpr.bold = True
        rpr.font.size = Pt(9.5)
        rpr.font.color.rgb = GREY
        for status, desc in ep["responses"]:
            b = doc.add_paragraph(style="List Bullet")
            sr = b.add_run(f"{status}: ")
            sr.bold = True
            sr.font.size = Pt(9)
            dr = b.add_run(desc)
            dr.font.size = Pt(9)

total = sum(len(e) for _, _, e in MODULES)
os.makedirs("docs", exist_ok=True)
out = os.path.join("docs", "API_Documentation.docx")
doc.save(out)
print("Wrote", out, "with", total, "endpoints across", len(MODULES), "modules")

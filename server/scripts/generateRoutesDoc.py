# Generates docx/API_Routes.docx documenting every current Express route.
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MAROON = RGBColor(0xB7, 0x1C, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)

# module -> (mount prefix, [ (method, path, access, validation, description) ])
MODULES = [
    ("Authentication", "/api/v1/auth", [
        ("POST", "/login", "Public", "loginSchema",
         "Log in with email + password. From a new/untrusted device it sends an email OTP and returns a challenge instead of a session. Rate limit: loginLimiter."),
        ("POST", "/verify_login_otp", "Public", "otpSchema",
         "Verify the emailed OTP for a new-device login, trust the device and issue the session. Rate limit: otpLimiter."),
        ("POST", "/create_admin", "Auth + Session (superadmin)", "adminCreationSchema",
         "Create a new admin account and email them their credentials. Superadmin only (enforced in controller)."),
        ("POST", "/refresh_token", "Public (refresh cookie)", "-",
         "Rotate the access + refresh tokens using the refresh-token cookie."),
        ("PATCH", "/update_password", "Auth + Session", "passwordSchema",
         "Change the logged-in admin's own password (verifies current password)."),
        ("GET", "/admins", "Auth + Session (superadmin)", "-",
         "List all admins. Superadmin only (enforced in controller)."),
        ("GET", "/profile/:id", "Auth", "-",
         "Get an admin profile. Allowed for the admin themselves or a superadmin."),
        ("POST", "/logout", "Auth", "-",
         "Log out: clears the stored session + refresh token and the auth cookies."),
        ("DELETE", "/remove_admin", "Auth + Session (superadmin)", "removeAdminSchema",
         "Delete an admin (cannot delete self or another superadmin). Superadmin only."),
    ]),
    ("Addresses", "/api/v1/addresses", [
        ("POST", "/", "Auth + Session", "addressSchema",
         "Create / set the current admin's address."),
        ("PATCH", "/", "Auth + Session", "addressSchema",
         "Update the current admin's address (same handler as POST)."),
        ("GET", "/", "Auth", "-",
         "Fetch the current admin's address."),
    ]),
    ("Media", "/api/v1/media", [
        ("POST", "/galleryUpload", "Auth", "galleryPostSchema",
         "Upload an image to the gallery (multipart 'file')."),
        ("DELETE", "/gallery/:id", "Auth", "-",
         "Delete a gallery item by id (also removes it from Cloudinary)."),
        ("GET", "/gallery", "Public", "galleryGetSchema",
         "List gallery items."),
        ("POST", "/mediaCoverage", "Auth", "mediaSchema",
         "Create a media-coverage entry (multipart 'file')."),
        ("DELETE", "/mediaCoverage", "Auth + Session", "-",
         "Delete a media-coverage entry."),
        ("GET", "/mediaCoverage/:id", "Public", "-",
         "Get a single media-coverage entry by id."),
        ("GET", "/mediaCoverage", "Public", "-",
         "List all media-coverage entries."),
    ]),
    ("Payments", "/api/v1/payments", [
        ("POST", "/webhook", "Public (Razorpay signature)", "-",
         "Razorpay server-to-server webhook. Authenticity verified via the webhook HMAC signature over the raw body; reconciles payment status."),
        ("POST", "/order", "Public", "createOrderSchema",
         "Create a standalone Razorpay order + Payment record. Domain flows (e.g. prasad booking) call the service directly instead."),
        ("POST", "/verify", "Public", "verifyPaymentSchema",
         "Verify a Razorpay payment signature and settle the payment."),
        ("GET", "/all", "Auth + Session", "-",
         "List all payments. Declared before '/:id' so the literal path is matched first."),
        ("GET", "/:id", "Auth + Session", "-",
         "Get a payment by its id."),
    ]),
    ("Prasad & Bookings", "/api/v1/prasad", [
        ("POST", "/add", "Auth + Session", "addPrasadSchema",
         "Add a prasad to the catalogue (multipart 'file' image)."),
        ("DELETE", "/remove/:id", "Auth + Session", "-",
         "Remove a prasad from the catalogue by id."),
        ("GET", "/prasads", "Public", "-",
         "List the full prasad catalogue."),
        ("GET", "/bookings", "Auth + Session", "-",
         "List all prasad bookings, newest first. Filters: ?status=pending|confirmed|failed, ?from=YYYY-MM-DD&to=YYYY-MM-DD, ?page & ?limit (default 20, max 100). Returns { bookings, pagination }."),
        ("POST", "/book", "Public", "bookPrasadSchema",
         "Create a prasad booking; server prices it and returns a Razorpay order."),
        ("POST", "/verify", "Public", "verifyPrasadBookingSchema",
         "Verify the Razorpay payment, confirm the booking, and email the receipt + PDF invoice."),
    ]),
]

METHOD_COLORS = {
    "GET": RGBColor(0x2E, 0x7D, 0x32),
    "POST": RGBColor(0x15, 0x65, 0xC0),
    "PATCH": RGBColor(0xE6, 0x5A, 0x00),
    "DELETE": RGBColor(0xC6, 0x28, 0x28),
}


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


doc = Document()

# Base document font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Shree Gaurishankar Baikunthdham Temple")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = MAROON

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Backend API - Route Reference")
r.font.size = Pt(13)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Base URL: /api/v1     |     Generated: {date.today().isoformat()}").italic = True

# --- Legend ---
doc.add_paragraph()
h = doc.add_paragraph().add_run("Access levels")
h.bold = True
h.font.size = Pt(12)
h.font.color.rgb = MAROON
for label, desc in [
    ("Public", "no authentication required."),
    ("Auth", "requires a valid access token (authenticate middleware)."),
    ("Auth + Session", "requires a valid access token AND an active session (authenticate + requireValidSession)."),
    ("(superadmin)", "additionally restricted to superadmin role, enforced inside the controller."),
]:
    p = doc.add_paragraph(style="List Bullet")
    br = p.add_run(f"{label}: ")
    br.bold = True
    p.add_run(desc)

note = doc.add_paragraph()
note.add_run("Note: ").bold = True
note.add_run("Most routes are throttled by apiRateLimiter; the auth login/OTP routes use dedicated stricter limiters. All responses use the standard ApiResponse envelope { statusCode, data, message, success }.")

# --- Per-module tables ---
COL_WIDTHS = [Inches(0.75), Inches(1.55), Inches(1.55), Inches(1.35), Inches(2.55)]
HEADERS = ["Method", "Path", "Access", "Validation", "Description"]

for name, prefix, routes in MODULES:
    doc.add_paragraph()
    hp = doc.add_paragraph()
    hr = hp.add_run(f"{name}   ({prefix})")
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = MAROON

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, text in enumerate(HEADERS):
        shade_cell(hdr[i], "B71C1C")
        para = hdr[i].paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)

    for method, path, access, validation, desc in routes:
        cells = table.add_row().cells
        mrun = cells[0].paragraphs[0].add_run(method)
        mrun.bold = True
        mrun.font.color.rgb = METHOD_COLORS.get(method, RGBColor(0, 0, 0))
        mrun.font.size = Pt(9)
        cells[1].paragraphs[0].add_run(prefix + path).font.size = Pt(9)
        cells[2].paragraphs[0].add_run(access).font.size = Pt(9)
        cells[3].paragraphs[0].add_run(validation).font.size = Pt(9)
        cells[4].paragraphs[0].add_run(desc).font.size = Pt(9)

    set_widths(table, COL_WIDTHS)

total = sum(len(r) for _, _, r in MODULES)
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run(f"{total} routes across {len(MODULES)} modules.")
fr.italic = True
fr.font.color.rgb = GREY

import os
os.makedirs("docx", exist_ok=True)
out = os.path.join("docx", "API_Routes.docx")
doc.save(out)
print("Wrote", out, "with", total, "routes")
